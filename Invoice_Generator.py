from docx import Document
class Invoice:
    def __init__(t, invoice_num, date, ddate, bto, items, trate):
        t.invoice_num = invoice_num
        t.date = date
        t.ddate = ddate
        t.bto = bto
        t.items = items
        t.trate = trate

    def ctotal(t):
        total = sum(item['quantity'] * item['unit_price'] for item in t.items)
        tax = total * t.trate
        total_with_tax = total + tax
        return total, tax, total_with_tax

    def generate_indata(t):
        total, tax, total_with_tax = t.ctotal()
        invoice_data = {
            'invoice_num': t.invoice_num,
            'date': t.date,
            'ddate': t.ddate,
            'bto': t.bto,
            'items': t.items,
            'total': total,
            'tax': tax,
            'total_with_tax': total_with_tax,
        }
        return invoice_data
def get_user_input():
    while True:
        try:
            invoice_num = input("Enter invoice ID: ")
            date = input("Date: ")
            ddate = input(" Due date: ")
            bto = input("This Bill Is To: ")
            items = []

            while True:
                description = input("Item description (or 'o' to stop): ")
                if description.lower() == 'o':
                    break
                quantity = int(input("Quantity: "))
                unit_price = float(input("Price: "))
                items.append({"description": description, "quantity": quantity, "unit_price": unit_price})

            trate = float(input("Tax rate % : "))
            
            if trate < 0 or any(item['quantity'] < 0 or item['unit_price'] < 0 for item in items):
                raise ValueError("Quantity, unit price, and tax rate must be non-negative.")
            
            return invoice_num, date, ddate, bto, items, trate
        
        except ValueError as e:
            print(f"Invalid input: {e}. Please try again.")
def generate_invoice(invoice_data):
    doc = Document()

    doc.add_heading('Invoice', 0)

    doc.add_paragraph(f"Invoice ID: {invoice_data['invoice_num']}")
    doc.add_paragraph(f"Date: {invoice_data['date']}")
    doc.add_paragraph(f"Due Date: {invoice_data['ddate']}\n")

    doc.add_paragraph(f"Bill To:\n{invoice_data['bto']}\n")

    table = doc.add_table(rows=1, cols=4)
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Description'
    hdr_cells[1].text = 'Quantity'
    hdr_cells[2].text = 'Unit Price'
    hdr_cells[3].text = 'Total'

    for item in invoice_data['items']:
        row_cells = table.add_row().cells
        row_cells[0].text = item['description']
        row_cells[1].text = str(item['quantity'])
        row_cells[2].text = f"${item['unit_price']:.2f}"
        row_cells[3].text = f"${item['quantity'] * item['unit_price']:.2f}"

    doc.add_paragraph(f"\n{'Total':<40} ${invoice_data['total']:.2f}")
    doc.add_paragraph(f"{'Tax':<40} ${invoice_data['tax']:.2f}")
    doc.add_paragraph(f"{'Total with Tax':<40} ${invoice_data['total_with_tax']:.2f}")

    doc.save("invoice.docx")
    print("Invoice saved as 'invoice.docx'")
def main():
    print("InsTextile welcomes you.... Kindly fill your details!!!:")

    while True:
        invoice_num, date, ddate, bto, items, trate = get_user_input()

        while True:
            invoice = Invoice(invoice_num, date, ddate, bto, items, trate)
            invoice_data = invoice.generate_indata()

            print("\nInvoice Details:")
            print(f"Invoice ID: {invoice_data['invoice_num']}")
            print(f"Date: {invoice_data['date']}")
            print(f"Due Date: {invoice_data['ddate']}")
            print(f"Bill To: {invoice_data['bto']}")
            print("Items:")
            for item in invoice_data['items']:
                print(f"{item['description']:30} {item['quantity']:>3} @ ${item['unit_price']:>6.2f} = ${item['quantity'] * item['unit_price']:>7.2f}")
            print(f"{'Total':40} ${invoice_data['total']:.2f}")
            print(f"{'Tax':40} ${invoice_data['tax']:.2f}")
            print(f"{'Total with Tax':40} ${invoice_data['total_with_tax']:.2f}")

            print("\nOptions:")
            print("1. Edit Invoice ID")
            print("2. Edit Date")
            print("3. Edit Due Date")
            print("4. Edit Bill To")
            print("5. Edit Items")
            print("6. Edit Tax Rate")
            print("7. Stop and Print Invoice")
            option = input("Choose an option (1-7): ")

            if option == '1':
                invoice_num = input("New invoice ID: ")
            elif option == '2':
                date = input("New invoice date: ")
            elif option == '3':
                ddate = input("New due date: ")
            elif option == '4':
                bto = input("New Bill info: ")
            elif option == '5':
                items = []
                while True:
                    description = input("Enter item description (or 'over' to stop): ")
                    if description.lower() == 'o':
                        break
                    quantity = int(input("Quantity: "))
                    unit_price = float(input("Price: "))
                    items.append({"description": description, "quantity": quantity, "unit_price": unit_price})
            elif option == '6':
                trate = float(input("New tax rate %: "))
            elif option == '7':
                generate_invoice(invoice_data)
                return
            else:
                print("Invalid option. Please try again.")

if __name__ == "__main__":
    main()
